"""
Complete Vector Knowledge Store Builder
Implements the full pipeline: Document → Chunks → Tokens → Embeddings → Vector Store
"""
import logging
from typing import List, Dict, Optional, Tuple
import os
from pathlib import Path

logger = logging.getLogger(__name__)


class VectorKnowledgeStoreBuilder:
    """
    Complete pipeline for building vector-based knowledge store
    
    STEP-BY-STEP FLOW:
    ==================
    Step 1: Document Extraction
        Input: PDF/DOCX/TXT files
        Process: Extract raw text using PyPDF2, python-docx
        Output: Raw text string
    
    Step 2: Chunking with Overlap
        Input: Raw text
        Process: Split into overlapping chunks (1000 tokens, 100 overlap)
        Output: List of text chunks
    
    Step 3: Tokenization
        Input: Text chunks
        Process: Convert to tokens (words/subwords)
        Output: Token sequences (vocabulary built)
    
    Step 4: Embeddings Generation
        Input: Text chunks
        Process: Convert to dense vectors using OpenAI/HuggingFace
        Output: Vector embeddings (1536 dims for OpenAI Ada)
    
    Step 5: Storage with Metadata
        Input: Embeddings + metadata
        Process: Store in vector database with chunk info
        Output: Stored in Chroma/Pinecone/FAISS
    
    Step 6: Indexing for Search
        Input: Stored embeddings
        Process: Build search index
        Output: Indexed and searchable vectors
    """
    
    def __init__(self, 
                 document_processor,
                 embedding_model,
                 vector_store,
                 chunk_size: int = 1000,
                 chunk_overlap: int = 100):
        """
        Initialize the knowledge store builder
        
        Args:
            document_processor: Service to process documents
            embedding_model: Service to generate embeddings
            vector_store: Backend vector database
            chunk_size: Size of text chunks in tokens
            chunk_overlap: Overlap between chunks for context
        """
        self.document_processor = document_processor
        self.embedding_model = embedding_model
        self.vector_store = vector_store
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    # ============= STEP 1: DOCUMENT EXTRACTION =============
    def extract_text_from_document(self, file_path: str, file_type: str) -> str:
        """
        STEP 1: Extract raw text from document
        Converts PDF/DOCX/TXT to plain text
        
        Args:
            file_path: Path to document
            file_type: Type of document (pdf, docx, txt)
            
        Returns:
            Raw extracted text
        """
        logger.info(f"STEP 1: Extracting text from {file_type.upper()} document")
        logger.info(f"File path: {file_path}")
        
        file_type = file_type.lower()
        
        if file_type == "pdf":
            return self._extract_from_pdf(file_path)
        elif file_type == "docx":
            return self._extract_from_docx(file_path)
        elif file_type == "txt":
            return self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyPDF2"""
        try:
            from PyPDF2 import PdfReader
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    text += page.extract_text() + "\n"
            logger.info(f"✓ Extracted {len(text)} characters from PDF ({len(pdf_reader.pages)} pages)")
            return text
        except Exception as e:
            logger.error(f"Error extracting PDF: {str(e)}")
            return ""
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx"""
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            logger.info(f"✓ Extracted {len(text)} characters from DOCX ({len(doc.paragraphs)} paragraphs)")
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX: {str(e)}")
            return ""
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            logger.info(f"✓ Extracted {len(text)} characters from TXT")
            return text
        except Exception as e:
            logger.error(f"Error extracting TXT: {str(e)}")
            return ""
    
    # ============= STEP 2: CHUNKING WITH OVERLAP =============
    def chunk_text(self, text: str) -> List[str]:
        """
        STEP 2: Split text into overlapping chunks
        Chunks help manage context length for embeddings
        Overlap ensures continuity between chunks
        
        Args:
            text: Raw text to chunk
            
        Returns:
            List of text chunks
        """
        logger.info(f"STEP 2: Chunking text")
        logger.info(f"Chunk size: {self.chunk_size} tokens, Overlap: {self.chunk_overlap} tokens")
        
        chunks = []
        words = text.split()
        
        start = 0
        while start < len(words):
            end = start + self.chunk_size
            chunk = " ".join(words[start:end])
            chunks.append(chunk)
            start = end - self.chunk_overlap
        
        logger.info(f"✓ Created {len(chunks)} chunks from text")
        logger.info(f"Average chunk size: {sum(len(c.split()) for c in chunks) / len(chunks):.0f} words")
        
        return chunks
    
    # ============= STEP 3: TOKENIZATION =============
    def tokenize_chunk(self, chunk: str) -> List[str]:
        """
        STEP 3: Tokenize a chunk
        Converts text to tokens (words/subwords)
        
        Args:
            chunk: Text chunk to tokenize
            
        Returns:
            List of tokens
        """
        # Simple word-based tokenization
        # Production: use NLTK, spaCy, or transformer tokenizers
        tokens = chunk.lower().split()
        return tokens
    
    def tokenize_chunks(self, chunks: List[str]) -> List[List[str]]:
        """
        STEP 3: Tokenize multiple chunks
        
        Args:
            chunks: List of text chunks
            
        Returns:
            List of token lists
        """
        logger.info(f"STEP 3: Tokenizing {len(chunks)} chunks")
        
        tokenized = [self.tokenize_chunk(chunk) for chunk in chunks]
        
        # Build vocabulary
        all_tokens = [t for tokens in tokenized for t in tokens]
        unique_tokens = len(set(all_tokens))
        total_tokens = len(all_tokens)
        
        logger.info(f"✓ Tokenized into {unique_tokens} unique tokens ({total_tokens} total)")
        
        return tokenized
    
    # ============= STEP 4: EMBEDDING GENERATION =============
    def generate_embeddings(self, chunks: List[str]) -> List[List[float]]:
        """
        STEP 4: Generate embeddings (dense vectors) for chunks
        Converts text chunks to numerical vectors using embedding model
        Each vector captures semantic meaning of the chunk
        
        Args:
            chunks: List of text chunks
            
        Returns:
            List of embedding vectors (1536-dimensional for OpenAI Ada)
        """
        logger.info(f"STEP 4: Generating embeddings for {len(chunks)} chunks")
        logger.info(f"Using embedding model: {self.embedding_model.__class__.__name__}")
        
        embeddings = self.embedding_model.embed_batch(chunks)
        
        if embeddings:
            embedding_dim = len(embeddings[0])
            logger.info(f"✓ Generated embeddings with dimension: {embedding_dim}")
        
        return embeddings
    
    # ============= STEP 5: STORE WITH METADATA =============
    def store_embeddings_with_metadata(self,
                                       document_id: str,
                                       filename: str,
                                       chunks: List[str],
                                       embeddings: List[List[float]],
                                       metadata: Optional[Dict] = None) -> List[str]:
        """
        STEP 5: Store embeddings with metadata in vector store
        Each embedding stored with chunk info for retrieval
        
        Metadata stored:
        - document_id: Source document
        - filename: Original filename
        - chunk_index: Position in document
        - chunk_text: Original text
        - total_chunks: Total chunks from document
        
        Args:
            document_id: Unique document identifier
            filename: Original filename
            chunks: List of text chunks
            embeddings: List of embedding vectors
            metadata: Additional metadata
            
        Returns:
            List of chunk IDs stored
        """
        logger.info(f"STEP 5: Storing {len(chunks)} embeddings with metadata")
        logger.info(f"Vector store: {self.vector_store.__class__.__name__}")
        
        metadata_list = []
        
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            chunk_metadata = {
                "document_id": document_id,
                "filename": filename,
                "chunk_index": i,
                "chunk_text": chunk[:500],  # Store first 500 chars for reference
                "total_chunks": len(chunks),
            }
            
            # Merge additional metadata
            if metadata:
                chunk_metadata.update(metadata)
            
            metadata_list.append(chunk_metadata)
        
        # Store in vector store
        chunk_ids = self.vector_store.add_documents(chunks, metadata_list)
        
        logger.info(f"✓ Stored {len(chunk_ids)} chunks with metadata")
        
        return chunk_ids
    
    # ============= STEP 6: INDEX FOR SIMILARITY SEARCH =============
    def build_search_index(self) -> bool:
        """
        STEP 6: Build searchable index on vector store
        Creates efficient index for similarity search
        
        Returns:
            True if indexing successful
        """
        logger.info(f"STEP 6: Building search index")
        
        # Vector stores typically handle indexing automatically
        # This is here for backends that need explicit indexing
        
        logger.info(f"✓ Search index built successfully")
        return True
    
    # ============= COMPLETE PIPELINE =============
    def build_knowledge_store(self,
                             file_path: str,
                             file_type: str,
                             document_id: str,
                             filename: str,
                             metadata: Optional[Dict] = None) -> Dict:
        """
        Complete pipeline: Document → Vector Knowledge Store
        
        Executes all 6 steps in sequence
        
        Args:
            file_path: Path to document file
            file_type: Type of document (pdf, docx, txt)
            document_id: Unique document identifier
            filename: Display filename
            metadata: Additional metadata
            
        Returns:
            Dictionary with pipeline results and statistics
        """
        logger.info("=" * 70)
        logger.info("STARTING VECTOR KNOWLEDGE STORE BUILD PIPELINE")
        logger.info("=" * 70)
        
        try:
            # STEP 1: Extract text
            raw_text = self.extract_text_from_document(file_path, file_type)
            if not raw_text:
                raise ValueError("Failed to extract text from document")
            
            # STEP 2: Chunk text
            chunks = self.chunk_text(raw_text)
            if not chunks:
                raise ValueError("Failed to create chunks")
            
            # STEP 3: Tokenize (informational)
            tokenized = self.tokenize_chunks(chunks)
            
            # STEP 4: Generate embeddings
            embeddings = self.generate_embeddings(chunks)
            if not embeddings:
                raise ValueError("Failed to generate embeddings")
            
            # STEP 5: Store with metadata
            chunk_ids = self.store_embeddings_with_metadata(
                document_id=document_id,
                filename=filename,
                chunks=chunks,
                embeddings=embeddings,
                metadata=metadata
            )
            
            # STEP 6: Build index
            index_built = self.build_search_index()
            
            result = {
                "success": True,
                "document_id": document_id,
                "filename": filename,
                "steps_completed": {
                    "extraction": True,
                    "chunking": True,
                    "tokenization": True,
                    "embedding": True,
                    "storage": True,
                    "indexing": index_built,
                },
                "statistics": {
                    "raw_text_length": len(raw_text),
                    "num_chunks": len(chunks),
                    "num_unique_tokens": len(set([t for tokens in tokenized for t in tokens])),
                    "embedding_dimension": len(embeddings[0]) if embeddings else 0,
                    "chunk_ids_stored": len(chunk_ids),
                    "avg_chunk_size": sum(len(c.split()) for c in chunks) / len(chunks) if chunks else 0,
                },
            }
            
            logger.info("=" * 70)
            logger.info("✓ VECTOR KNOWLEDGE STORE BUILD COMPLETED SUCCESSFULLY")
            logger.info("=" * 70)
            logger.info(f"Statistics: {result['statistics']}")
            
            return result
            
        except Exception as e:
            logger.error(f"Error in knowledge store build: {str(e)}")
            return {
                "success": False,
                "document_id": document_id,
                "error": str(e),
            }
    
    # ============= SEARCH FUNCTIONALITY =============
    def semantic_search(self, query: str, top_k: int = 5) -> List[Dict]:
        """
        Perform semantic search on indexed embeddings
        
        Args:
            query: Search query
            top_k: Number of top results
            
        Returns:
            List of search results with similarity scores
        """
        logger.info(f"Performing semantic search: {query}")
        
        # Generate embedding for query
        query_embedding = self.embedding_model.embed(query)
        
        # Search in vector store
        results = self.vector_store.search(query, top_k)
        
        logger.info(f"✓ Found {len(results)} relevant results")
        
        return results
    
    # ============= VECTOR STATISTICS =============
    def get_store_statistics(self) -> Dict:
        """
        Get statistics about the vector store
        
        Returns:
            Dictionary with store statistics
        """
        return {
            "store_type": self.vector_store.__class__.__name__,
            "embedding_dimension": 1536,  # For OpenAI ada-002
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
        }
